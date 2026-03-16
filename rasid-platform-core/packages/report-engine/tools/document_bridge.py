import argparse
import base64
import json
import re
import unicodedata
import zipfile
from pathlib import Path

import fitz
import arabic_reshaper
from bidi.algorithm import get_display
from docx import Document
from docx.document import Document as _Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph
from reportlab.lib import colors
from reportlab.lib.enums import TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.graphics.shapes import Drawing, Line, Rect, String
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph as PdfParagraph, SimpleDocTemplate, Spacer, Table as PdfTable, TableStyle

CAPTION_PREFIX_RE = re.compile(r"^(table|figure|chart|شكل|جدول)\s*([\d١٢٣٤٥٦٧٨٩٠]+)[:\-]?\s*(.+)$", re.IGNORECASE)
CAPTION_SUFFIX_RE = re.compile(
    r"^(?:(?P<leading>[\d١٢٣٤٥٦٧٨٩٠\s.,%/-]+)\s*[:\-]\s*)?(?P<body>.+?)\s+(?P<number>[\d١٢٣٤٥٦٧٨٩٠]+)\s+(?P<label>table|figure|chart|شكل|جدول)$",
    re.IGNORECASE
)
HEADING_RE = re.compile(r"^(executive summary|summary|appendix|section|overview|مقدمة|ملخص|الملخص|الملحق|القسم)\b", re.IGNORECASE)
ARABIC_RE = re.compile(r"[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]")
NUMBER_RE = re.compile(r"[-+]?\d+(?:\.\d+)?")
SAMPLE_PNG_BASE64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+nm9kAAAAASUVORK5CYII="
)

def bbox(x, y, w, h, unit="pt"):
    return {"x": round(float(x), 2), "y": round(float(y), 2), "w": round(float(w), 2), "h": round(float(h), 2), "unit": unit}

def normalize_title(text):
    return re.sub(r"\s+", " ", (text or "").strip())

def normalize_source_text(text):
    value = unicodedata.normalize("NFKC", text or "")
    value = re.sub(r"([\u0600-\u06FFA-Za-z])([\d١٢٣٤٥٦٧٨٩٠])", r"\1 \2", value)
    value = re.sub(r"([\d١٢٣٤٥٦٧٨٩٠])([\u0600-\u06FFA-Za-z])", r"\1 \2", value)
    return normalize_title(value)

def shape(v, rtl):
    return get_display(arabic_reshaper.reshape(v)) if rtl and v else v or ""

def ar_font():
    name = "ArialArabic"
    candidates = [
        r"C:\Windows\Fonts\arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
        "/usr/share/fonts/TTF/DejaVuSans.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
    ]
    try:
        pdfmetrics.getFont(name)
    except Exception:
        font_path = next((c for c in candidates if Path(c).exists()), None)
        if font_path:
            pdfmetrics.registerFont(TTFont(name, font_path))
        else:
            return "Helvetica"
    return name

def rtl_ratio(v):
    return len(ARABIC_RE.findall(v or "")) / max(1, len(v or ""))

def infer_lang(v):
    return "ar" if ARABIC_RE.search(v or "") else "en"

def section_kind(title, index):
    low = (title or "").lower()
    if index == 0:
        return "cover"
    if "appendix" in low or "ملحق" in title:
        return "appendix"
    if "summary" in low or "ملخص" in title:
        return "executive_summary"
    return "body"

def parse_caption(text):
    normalized = normalize_source_text(text)
    if not normalized:
        return None

    def caption_payload(label, body, number=None, leading=None):
        canonical_label = normalize_source_text(label)
        clean_body = re.sub(r"^[:\-]\s*", "", normalize_title(body or normalized))
        leading_values = [float(v) for v in NUMBER_RE.findall(leading or "")]
        body_values = [float(v) for v in NUMBER_RE.findall(clean_body)]
        return {
            "label": canonical_label,
            "text": clean_body,
            "target_kind": "table" if canonical_label.lower().startswith("table") or canonical_label == "جدول" else "chart",
            "caption_number": str(number) if number is not None else None,
            "canonical_text": f"{canonical_label} {number}: {clean_body}".strip() if number is not None else f"{canonical_label}: {clean_body}".strip(),
            "series_values": body_values[:6] if body_values else leading_values[:6]
        }

    prefix_match = CAPTION_PREFIX_RE.match(normalized)
    if prefix_match:
        return caption_payload(prefix_match.group(1), prefix_match.group(3), prefix_match.group(2))

    suffix_match = CAPTION_SUFFIX_RE.match(normalized)
    if suffix_match:
        return caption_payload(
            suffix_match.group("label"),
            suffix_match.group("body"),
            suffix_match.group("number"),
            suffix_match.group("leading")
        )

    lowered = normalized.lower()
    for prefix in ["figure", "table", "chart", "شكل", "جدول"]:
        if lowered.startswith(prefix):
            body = normalized.split(":", 1)[1].strip() if ":" in normalized else normalized[len(prefix):].strip()
            return caption_payload(prefix, body)
    return None

def alignment(paragraph):
    mapping = {0: "left", 1: "center", 2: "right", 3: "justify"}
    return mapping.get(int(paragraph.alignment), "inherit") if paragraph.alignment is not None else "inherit"

def iter_blocks(parent):
    body = parent.element.body if isinstance(parent, _Document) else None
    if body is None:
        raise ValueError("unsupported parent")
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def page_break(paragraph):
    xml = paragraph._element.xml
    return 'w:type="page"' in xml or "lastRenderedPageBreak" in xml

def serialize_docx_sections(document):
    sections = []
    for index, section in enumerate(document.sections):
        sections.append({
            "section_index": index + 1,
            "start_type": str(section.start_type) if section.start_type is not None else str(WD_SECTION_START.NEW_PAGE),
            "orientation": "landscape" if section.orientation == WD_ORIENT.LANDSCAPE else "portrait",
            "page_width_pt": round(float(section.page_width.pt), 2),
            "page_height_pt": round(float(section.page_height.pt), 2),
            "margin_top_pt": round(float(section.top_margin.pt), 2),
            "margin_bottom_pt": round(float(section.bottom_margin.pt), 2),
            "margin_left_pt": round(float(section.left_margin.pt), 2),
            "margin_right_pt": round(float(section.right_margin.pt), 2),
            "header_text": normalize_title(section.header.paragraphs[0].text if section.header.paragraphs else ""),
            "footer_text": normalize_title(section.footer.paragraphs[0].text if section.footer.paragraphs else "")
        })
    return sections

def docx_relationship_maps(document, input_path):
    hyperlinks = {}
    media_assets = []
    embedded_assets = []
    for rel_id, rel in document.part.rels.items():
        rel_type = getattr(rel, "reltype", "")
        target_ref = getattr(rel, "target_ref", None) or str(getattr(rel, "target_part", "") or "")
        if rel_type == RT.HYPERLINK:
            hyperlinks[rel_id] = target_ref
        elif rel_type == RT.IMAGE:
            media_assets.append({"asset_id": rel_id, "asset_kind": "image", "target_ref": target_ref})
        elif "oleObject" in rel_type or "package" in rel_type:
            embedded_assets.append({"asset_id": rel_id, "asset_kind": "embedded_package", "target_ref": target_ref})
    with zipfile.ZipFile(input_path, "r") as archive:
        for name in archive.namelist():
            if name.startswith("word/media/"):
                media_assets.append({"asset_id": f"media::{Path(name).name}", "asset_kind": "media_entry", "target_ref": name})
            elif name.startswith("word/embeddings/"):
                embedded_assets.append({"asset_id": f"embedding::{Path(name).name}", "asset_kind": "embedded_package", "target_ref": name})
    return hyperlinks, media_assets, embedded_assets

def paragraph_links(paragraph, hyperlink_targets):
    links = []
    for hyperlink in paragraph._element.xpath(".//w:hyperlink"):
        rel_id = hyperlink.get(qn("r:id"))
        text = "".join(node.text or "" for node in hyperlink.xpath(".//w:t"))
        links.append({
            "text": normalize_title(text),
            "rel_id": rel_id,
            "target": hyperlink_targets.get(rel_id)
        })
    return [link for link in links if link["text"] or link["target"]]

def paragraph_drawing_count(paragraph):
    return len(paragraph._element.xpath(".//w:drawing"))

def docx_core_metadata(document):
    props = document.core_properties
    return {
        "author": props.author,
        "category": props.category,
        "comments": props.comments,
        "created": props.created.isoformat() if props.created else None,
        "modified": props.modified.isoformat() if props.modified else None,
        "subject": props.subject,
        "title": props.title
    }

def add_hyperlink(paragraph, text, url):
    relation_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(color)
    properties.append(underline)
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)

def write_sample_png(output_dir):
    target = Path(output_dir) / "sample-chart.png"
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_bytes(base64.b64decode(SAMPLE_PNG_BASE64))
    return target

def extract_docx(input_path):
    document = Document(input_path)
    title = None
    extracted_lines, paragraphs, layout_semantics, page_structure = [], [], [], []
    captions, tables, charts, section_blueprints = [], [], [], []
    geometry_map, page_semantics, section_hierarchy, block_lineage_map = [], [], [], []
    hyperlink_targets, media_assets, embedded_assets = docx_relationship_maps(document, input_path)
    hyperlinks = []
    document_sections = serialize_docx_sections(document)
    current_section, current_heading_level, pending_caption, page_number, page_blocks = None, None, None, 1, 0
    preface_blocks = []
    page_heading_count = 0
    page_caption_count = 0
    page_table_count = 0
    page_chart_count = 0

    def flush_page():
        nonlocal page_blocks, page_heading_count, page_caption_count, page_table_count, page_chart_count
        page_structure.append({
            "page_number": len(page_structure) + 1,
            "block_count": page_blocks,
            "text_direction": "rtl" if any(item.get("rtl") for item in layout_semantics if item.get("page_number") == len(page_structure) + 1) else "ltr",
            "heading_count": page_heading_count,
            "caption_count": page_caption_count,
            "table_count": page_table_count,
            "chart_count": page_chart_count,
            "hyperlink_count": len([link for link in hyperlinks if link.get("page_number") == len(page_structure) + 1]),
            "media_count": len(media_assets)
        })
        page_semantics.append({
            "page_number": len(page_semantics) + 1,
            "semantic_role": "report_page",
            "reading_direction": "rtl" if any(item.get("rtl") for item in layout_semantics if item.get("page_number") == len(page_semantics) + 1) else "ltr",
            "heading_count": page_heading_count,
            "caption_count": page_caption_count,
            "table_count": page_table_count,
            "chart_count": page_chart_count,
            "hyperlink_count": len([link for link in hyperlinks if link.get("page_number") == len(page_semantics) + 1]),
            "document_section": document_sections[min(len(page_semantics), max(0, len(document_sections) - 1))] if document_sections else None
        })
        page_blocks = 0
        page_heading_count = 0
        page_caption_count = 0
        page_table_count = 0
        page_chart_count = 0

    def ensure_section(name, level=1):
        nonlocal current_section, current_heading_level
        if current_section and current_section["title"] == name:
            return
        current_section = {"title": name, "section_kind": section_kind(name, len(section_blueprints)), "blocks": []}
        section_blueprints.append(current_section)
        current_heading_level = level
        section_hierarchy.append({
            "section_title": name,
            "heading_level": level,
            "page_number": page_number,
            "section_kind": current_section["section_kind"],
            "section_index": len(section_blueprints)
        })

    for block_index, block in enumerate(iter_blocks(document)):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if page_break(block):
                flush_page()
                page_number += 1
            if not text:
                continue
            title = title or text
            extracted_lines.append(text)
            page_blocks += 1
            style_name = block.style.name if block.style else "Normal"
            rtl = rtl_ratio(text) > 0.2
            heading_level = 1 if style_name.lower().startswith("heading 1") else 2 if style_name.lower().startswith("heading 2") else 3 if style_name.lower().startswith("heading 3") else 0
            links = paragraph_links(block, hyperlink_targets)
            drawing_count = paragraph_drawing_count(block)
            geometry = bbox(48, 72 + (page_blocks - 1) * 22, 499, 18 if heading_level else 14)
            layout_semantics.append({"kind": "paragraph", "page_number": page_number, "style_name": style_name, "alignment": alignment(block), "text": text, "rtl": rtl, "heading_level": heading_level, "bbox": geometry, "hyperlinks": links, "drawing_count": drawing_count})
            geometry_map.append({
                "geometry_id": f"docx-geometry-{len(geometry_map) + 1}",
                "page_number": page_number,
                "order_index": page_blocks,
                "kind": "paragraph",
                "bbox": geometry,
                "text": text,
                "heading_level": heading_level,
                "rtl": rtl,
                "hyperlink_count": len(links),
                "drawing_count": drawing_count
            })
            for link_index, link in enumerate(links):
                link_ref = f"docx://link/{len(hyperlinks) + 1}"
                hyperlinks.append({
                    "link_ref": link_ref,
                    "page_number": page_number,
                    "paragraph_ref": f"docx://paragraph/{len(extracted_lines)}",
                    "text": link["text"],
                    "target": link["target"],
                    "rel_id": link["rel_id"]
                })
                block_lineage_map.append({
                    "source_ref": link_ref,
                    "target_kind": "hyperlink",
                    "target_ref": link["target"] or link["text"],
                    "page_number": page_number,
                    "paragraph_ref": f"docx://paragraph/{len(extracted_lines)}",
                    "link_index": link_index + 1
                })
            caption = parse_caption(text)
            if caption:
                pending_caption = {"caption_id": f"caption-{len(captions) + 1}", "page_number": page_number, **caption}
                captions.append(pending_caption)
                page_caption_count += 1
                continue
            if style_name.lower() == "title":
                continue
            is_heading = style_name.lower().startswith("heading") or HEADING_RE.match(text)
            if is_heading or current_section is None:
                if is_heading:
                    ensure_section(text, max(1, heading_level or 1))
                    if preface_blocks:
                        current_section["blocks"].extend(preface_blocks)
                        preface_blocks = []
                elif current_section is None:
                    paragraphs.append(text)
                    lineage_ref = f"docx://paragraph/{len(extracted_lines)}"
                    preface_blocks.append({"block_type": "narrative", "title": text[:80], "body": text, "dataset_ref": None, "query_ref": None, "field_mappings": [], "citations": [], "table_rows": [], "chart_series": [], "source_metadata": {"page_number": page_number, "style_name": style_name, "section_heading_level": current_heading_level, "hyperlinks": links, "drawing_count": drawing_count, "document_section": document_sections[min(page_number - 1, max(0, len(document_sections) - 1))] if document_sections else None}, "layout_semantics": {"page_number": page_number, "alignment": alignment(block), "rtl": rtl, "bbox": geometry, "hyperlinks": links}, "caption": "", "page_number": page_number, "source_lineage_refs": [lineage_ref]})
                    block_lineage_map.append({"source_ref": lineage_ref, "target_kind": "paragraph", "target_ref": lineage_ref, "page_number": page_number, "section_title": "preface"})
                    continue
                if is_heading:
                    page_heading_count += 1
                    paragraphs.append(text)
                    continue
            nums = [float(v) for v in NUMBER_RE.findall(text)]
            if pending_caption and pending_caption["target_kind"] == "chart":
                series_values = nums[:6] if nums else pending_caption.get("series_values", [])[:6]
                chart_series = [{"label": f"Point {i + 1}", "value": v} for i, v in enumerate(series_values)]
                chart_id = f"chart-{len(charts) + 1}"
                charts.append({"chart_id": chart_id, "page_number": page_number, "title": pending_caption["text"], "caption": pending_caption["text"], "chart_type": "bar", "series": chart_series, "source_anchor": text, "geometry": geometry, "hyperlinks": links, "drawing_count": drawing_count})
                current_section["blocks"].append({"block_type": "chart", "title": pending_caption["text"], "body": text, "dataset_ref": None, "query_ref": None, "field_mappings": [], "citations": [pending_caption["caption_id"]], "table_rows": [], "chart_series": chart_series, "source_metadata": {"page_number": page_number, "style_name": style_name, "caption_id": pending_caption["caption_id"], "source_anchor": text, "section_heading_level": current_heading_level, "hyperlinks": links, "drawing_count": drawing_count}, "layout_semantics": {"page_number": page_number, "alignment": alignment(block), "rtl": rtl, "bbox": geometry, "page_semantics_role": "chart_caption_region", "document_section": document_sections[min(page_number - 1, max(0, len(document_sections) - 1))] if document_sections else None}, "caption": pending_caption["text"], "page_number": page_number, "source_lineage_refs": [f"docx://paragraph/{len(extracted_lines)}", f"docx://chart/{chart_id}"]})
                block_lineage_map.append({"source_ref": f"docx://paragraph/{len(extracted_lines)}", "target_kind": "chart", "target_ref": chart_id, "caption_ref": pending_caption["caption_id"], "page_number": page_number})
                page_chart_count += 1
                pending_caption = None
                continue
            paragraphs.append(text)
            lineage_ref = f"docx://paragraph/{len(extracted_lines)}"
            current_section["blocks"].append({"block_type": "narrative", "title": text[:80], "body": text, "dataset_ref": None, "query_ref": None, "field_mappings": [], "citations": [], "table_rows": [], "chart_series": [], "source_metadata": {"page_number": page_number, "style_name": style_name, "section_heading_level": current_heading_level, "hyperlinks": links, "drawing_count": drawing_count, "document_section": document_sections[min(page_number - 1, max(0, len(document_sections) - 1))] if document_sections else None}, "layout_semantics": {"page_number": page_number, "alignment": alignment(block), "rtl": rtl, "bbox": geometry, "hyperlinks": links}, "caption": "", "page_number": page_number, "source_lineage_refs": [lineage_ref]})
            block_lineage_map.append({"source_ref": lineage_ref, "target_kind": "paragraph", "target_ref": lineage_ref, "page_number": page_number, "section_title": current_section["title"]})
        else:
            page_blocks += 1
            rows = [[cell.text.strip() for cell in row.cells] for row in block.rows]
            table_caption = pending_caption["text"] if pending_caption and pending_caption["target_kind"] == "table" else f"Imported Table {len(tables) + 1}"
            table_id = f"table-{len(tables) + 1}"
            table_geometry = bbox(48, 72 + (page_blocks - 1) * 22, 499, max(40, len(rows) * 18))
            table_record = {"table_id": table_id, "page_number": page_number, "title": table_caption, "caption": table_caption, "rows": rows, "row_count": len(rows), "column_count": len(rows[0]) if rows and rows[0] else 0, "geometry": table_geometry}
            tables.append(table_record)
            layout_semantics.append({"kind": "table", "page_number": page_number, "title": table_caption, "row_count": table_record["row_count"], "column_count": table_record["column_count"], "bbox": table_geometry})
            geometry_map.append({
                "geometry_id": f"docx-geometry-{len(geometry_map) + 1}",
                "page_number": page_number,
                "order_index": page_blocks,
                "kind": "table",
                "bbox": table_geometry,
                "table_id": table_id,
                "caption": table_caption,
                "row_headers": rows[0] if rows else []
            })
            ensure_section(current_section["title"] if current_section else "Imported Tables")
            current_section["blocks"].append({"block_type": "table", "title": table_caption, "body": table_caption, "dataset_ref": None, "query_ref": None, "field_mappings": [], "citations": [pending_caption["caption_id"]] if pending_caption else [], "table_rows": rows, "chart_series": [], "source_metadata": {"page_number": page_number, "row_count": table_record["row_count"], "column_count": table_record["column_count"], "section_heading_level": current_heading_level, "document_section": document_sections[min(page_number - 1, max(0, len(document_sections) - 1))] if document_sections else None}, "layout_semantics": {"page_number": page_number, "kind": "table", "bbox": table_geometry, "page_semantics_role": "table_region"}, "caption": table_caption, "page_number": page_number, "source_lineage_refs": [f"docx://table/{len(tables)}"]})
            block_lineage_map.append({"source_ref": f"docx://table/{len(tables)}", "target_kind": "table", "target_ref": table_id, "caption_ref": pending_caption["caption_id"] if pending_caption else None, "page_number": page_number})
            page_table_count += 1
            pending_caption = None

    flush_page()
    section_blueprints = [section for section in section_blueprints if section["blocks"]]
    if preface_blocks and section_blueprints:
        section_blueprints[0]["blocks"] = [*preface_blocks, *section_blueprints[0]["blocks"]]
    if not section_blueprints:
        ensure_section(title or Path(input_path).stem)
        if preface_blocks:
            current_section["blocks"].extend(preface_blocks)
        else:
            fallback_ref = f"docx://fallback/{Path(input_path).stem}"
            current_section["blocks"].append({"block_type": "narrative", "title": title or Path(input_path).stem, "body": "\n".join(paragraphs) or Path(input_path).stem, "dataset_ref": None, "query_ref": None, "field_mappings": [], "citations": [], "table_rows": [], "chart_series": [], "source_metadata": {"page_number": 1, "fallback": True}, "layout_semantics": {"page_number": 1, "kind": "fallback", "bbox": bbox(48, 72, 499, 40)}, "caption": "", "page_number": 1, "source_lineage_refs": [fallback_ref]})
            block_lineage_map.append({"source_ref": fallback_ref, "target_kind": "paragraph", "target_ref": fallback_ref, "page_number": 1, "section_title": title or Path(input_path).stem})
    return {"title": title or Path(input_path).stem, "paragraphs": paragraphs, "extracted_text": "\n".join(extracted_lines), "warning_codes": [] if extracted_lines else ["docx_empty_extract"], "page_count": max(1, len(page_structure)), "section_count": len(section_blueprints), "table_count": len(tables), "chart_count": len(charts), "caption_count": len(captions), "captions": captions, "tables": tables, "charts": charts, "page_structure": page_structure, "layout_semantics": layout_semantics, "geometry_map": geometry_map, "page_semantics": page_semantics, "section_hierarchy": section_hierarchy, "block_lineage_map": block_lineage_map, "section_blueprints": section_blueprints, "rendered_pages": [], "source_language": infer_lang("\n".join(extracted_lines)), "hyperlinks": hyperlinks, "embedded_assets": media_assets + embedded_assets, "document_metadata": {"core_properties": docx_core_metadata(document), "sections": document_sections}}

def pdf_direction(text):
    return "rtl" if rtl_ratio(text) > 0.15 else "ltr"

def extract_pdf(input_path, render_dir=None):
    doc = fitz.open(input_path)
    extracted_lines, paragraphs, layout_semantics, page_structure = [], [], [], []
    captions, tables, charts, rendered_pages, section_blueprints = [], [], [], [], []
    geometry_map, page_semantics, section_hierarchy, block_lineage_map = [], [], [], []
    hyperlinks, embedded_assets = [], []
    current_section, pending_caption, title = None, None, None
    current_heading_level = None
    preface_blocks = []
    pending_table_captions = []

    def ensure_section(name, level=1):
        nonlocal current_section, current_heading_level
        if current_section and current_section["title"] == name:
            return
        current_section = {"title": name, "section_kind": section_kind(name, len(section_blueprints)), "blocks": []}
        section_blueprints.append(current_section)
        current_heading_level = level
        section_hierarchy.append({
            "section_title": normalize_title(name),
            "heading_level": level,
            "page_number": max(1, len(page_structure) + 1),
            "section_kind": current_section["section_kind"],
            "section_index": len(section_blueprints)
        })

    for page_index, page in enumerate(doc):
        text_dict = page.get_text("dict", sort=True)
        text_blocks = []
        for block in text_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
            lines, max_font = [], 0
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    value = (span.get("text") or "").strip()
                    if value:
                        lines.append(value)
                    max_font = max(max_font, float(span.get("size") or 0))
            text = " ".join(lines).strip()
            if text:
                words = page.get_text("words")
                text_blocks.append({"text": text, "bbox": block.get("bbox"), "font_size": max_font, "direction": pdf_direction(text), "word_count": len([word for word in words if len(word) >= 5 and word[4] == block.get("number", word[4])])})
        drawings = page.get_drawings()
        image_count = len(page.get_images(full=True))
        page_links = page.get_links()
        try:
            page_images = page.get_image_info(xrefs=True)
        except Exception:
            page_images = []
        try:
            found_tables = getattr(page.find_tables(), "tables", [])
        except Exception:
            found_tables = []
        page_structure.append({"page_number": page_index + 1, "block_count": len(text_blocks) + len(found_tables) + image_count + len(drawings), "text_direction": "rtl" if any(item["direction"] == "rtl" for item in text_blocks) else "ltr", "image_count": image_count, "drawing_count": len(drawings), "table_count": len(found_tables), "chart_region_count": len([drawing for drawing in drawings if drawing.get("items")]), "dominant_font_size": max([item["font_size"] for item in text_blocks], default=0), "hyperlink_count": len(page_links), "vector_region_count": len(drawings)})
        page_semantics.append({
            "page_number": page_index + 1,
            "semantic_role": "report_page",
            "reading_direction": "rtl" if any(item["direction"] == "rtl" for item in text_blocks) else "ltr",
            "heading_count": len([item for item in text_blocks if item["font_size"] >= 15]),
            "table_count": len(found_tables),
            "chart_region_count": len(drawings) + image_count,
            "dense_table_page": len(found_tables) > 0 and len(text_blocks) < 10,
            "hyperlink_count": len(page_links),
            "vector_region_count": len(drawings)
        })
        if render_dir:
            target_dir = Path(render_dir)
            target_dir.mkdir(parents=True, exist_ok=True)
            png_path = target_dir / f"page-{page_index + 1}.png"
            page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False).save(str(png_path))
            rendered_pages.append(str(png_path))
        for block_order, entry in enumerate(text_blocks):
            text = entry["text"]
            title = title or text
            extracted_lines.append(text)
            bbox_value = {"x0": round(entry["bbox"][0], 2), "y0": round(entry["bbox"][1], 2), "x1": round(entry["bbox"][2], 2), "y1": round(entry["bbox"][3], 2), "unit": "pt"} if entry["bbox"] else None
            heading_level = 1 if entry["font_size"] >= 18 else 2 if entry["font_size"] >= 15 else 0
            normalized_text = normalize_source_text(text)
            related_links = []
            if bbox_value:
                for link in page_links:
                    rect = fitz.Rect(link.get("from", (0, 0, 0, 0)))
                    if rect.intersects(fitz.Rect(entry["bbox"])):
                        link_ref = f"pdf://page/{page_index + 1}/link/{len(hyperlinks) + 1}"
                        link_record = {"link_ref": link_ref, "page_number": page_index + 1, "text_block_ref": f"pdf://page/{page_index + 1}/text/{len(extracted_lines)}", "uri": link.get("uri"), "kind": link.get("kind"), "bbox": bbox(rect.x0, rect.y0, rect.width, rect.height)}
                        hyperlinks.append(link_record)
                        related_links.append(link_record)
            layout_semantics.append({"kind": "text_block", "page_number": page_index + 1, "bbox": bbox_value, "font_size": entry["font_size"], "direction": entry["direction"], "text": text, "heading_level": heading_level, "hyperlinks": related_links})
            geometry_map.append({
                "geometry_id": f"pdf-geometry-{len(geometry_map) + 1}",
                "page_number": page_index + 1,
                "order_index": block_order + 1,
                "kind": "text_block",
                "bbox": bbox_value,
                "font_size": entry["font_size"],
                "direction": entry["direction"],
                "text": text,
                "hyperlink_count": len(related_links)
            })
            is_margin_marker = bbox_value and (bbox_value["y0"] < 32 or bbox_value["y1"] > page.rect.height - 20)
            is_running_header = is_margin_marker and entry["font_size"] <= 10 and ("rasid" in normalized_text.lower() or normalized_text.lower().startswith("page "))
            is_document_title = page_index == 0 and heading_level == 1 and block_order <= 1
            if is_running_header or is_document_title:
                if is_document_title:
                    title = normalized_text
                continue
            caption = parse_caption(text)
            if caption:
                caption_record = {"caption_id": f"caption-{len(captions) + 1}", "page_number": page_index + 1, **caption}
                captions.append(caption_record)
                if caption["target_kind"] == "table":
                    pending_table_captions.append(caption_record)
                else:
                    pending_caption = caption_record
                continue
            is_heading = entry["font_size"] >= 15 or HEADING_RE.match(normalized_text)
            if is_heading or current_section is None:
                if is_heading:
                    ensure_section(normalized_text, max(1, heading_level or 1))
                    if preface_blocks:
                        current_section["blocks"].extend(preface_blocks)
                        preface_blocks = []
                elif current_section is None:
                    paragraphs.append(text)
                    lineage_ref = f"pdf://page/{page_index + 1}/text/{len(paragraphs)}"
                    preface_blocks.append({"block_type": "narrative", "title": text[:80], "body": text, "dataset_ref": None, "query_ref": None, "field_mappings": [], "citations": [], "table_rows": [], "chart_series": [], "source_metadata": {"page_number": page_index + 1, "bbox": bbox_value, "section_heading_level": current_heading_level, "hyperlinks": related_links}, "layout_semantics": {"page_number": page_index + 1, "font_size": entry["font_size"], "direction": entry["direction"], "bbox": bbox_value, "hyperlinks": related_links}, "caption": "", "page_number": page_index + 1, "source_lineage_refs": [lineage_ref]})
                    block_lineage_map.append({"source_ref": lineage_ref, "target_kind": "paragraph", "target_ref": lineage_ref, "page_number": page_index + 1, "section_title": "preface"})
                    continue
                if is_heading:
                    continue
            nums = [float(v) for v in NUMBER_RE.findall(text)]
            if pending_caption and pending_caption["target_kind"] == "chart":
                series_values = nums[:6] if nums else pending_caption.get("series_values", [])[:6]
                chart_signal_text = any(token in normalized_text.lower() for token in ["chart", "trend", "vector", "coverage"])
                if series_values or chart_signal_text:
                    chart_series = [{"label": f"Point {i + 1}", "value": v} for i, v in enumerate(series_values)]
                    chart_id = f"chart-{len(charts) + 1}"
                    charts.append({"chart_id": chart_id, "page_number": page_index + 1, "title": pending_caption["text"], "caption": pending_caption["text"], "chart_type": "bar", "series": chart_series, "source_anchor": text, "geometry": bbox_value})
                    current_section["blocks"].append({"block_type": "chart", "title": pending_caption["text"], "body": text, "dataset_ref": None, "query_ref": None, "field_mappings": [], "citations": [pending_caption["caption_id"]], "table_rows": [], "chart_series": chart_series, "source_metadata": {"page_number": page_index + 1, "bbox": bbox_value, "drawings_detected": len(drawings), "images_detected": image_count, "section_heading_level": current_heading_level, "hyperlinks": related_links}, "layout_semantics": {"page_number": page_index + 1, "font_size": entry["font_size"], "direction": entry["direction"], "bbox": bbox_value, "page_semantics_role": "chart_caption_region", "hyperlinks": related_links}, "caption": pending_caption["text"], "page_number": page_index + 1, "source_lineage_refs": [f"pdf://page/{page_index + 1}/chart/{len(charts) + 1}", chart_id]})
                    block_lineage_map.append({"source_ref": f"pdf://page/{page_index + 1}/chart/{len(charts) + 1}", "target_kind": "chart", "target_ref": chart_id, "caption_ref": pending_caption["caption_id"], "page_number": page_index + 1})
                    pending_caption = None
                    continue
            paragraphs.append(text)
            lineage_ref = f"pdf://page/{page_index + 1}/text/{len(paragraphs)}"
            narrative_block = {"block_type": "narrative", "title": text[:80], "body": text, "dataset_ref": None, "query_ref": None, "field_mappings": [], "citations": [], "table_rows": [], "chart_series": [], "source_metadata": {"page_number": page_index + 1, "bbox": bbox_value, "section_heading_level": current_heading_level, "hyperlinks": related_links}, "layout_semantics": {"page_number": page_index + 1, "font_size": entry["font_size"], "direction": entry["direction"], "bbox": bbox_value, "hyperlinks": related_links}, "caption": "", "page_number": page_index + 1, "source_lineage_refs": [lineage_ref]}
            if current_section is None:
                preface_blocks.append(narrative_block)
                block_lineage_map.append({"source_ref": lineage_ref, "target_kind": "paragraph", "target_ref": lineage_ref, "page_number": page_index + 1, "section_title": "preface"})
            else:
                current_section["blocks"].append(narrative_block)
                block_lineage_map.append({"source_ref": lineage_ref, "target_kind": "paragraph", "target_ref": lineage_ref, "page_number": page_index + 1, "section_title": current_section["title"]})
        for table_index, table in enumerate(found_tables):
            try:
                rows = [[("" if cell is None else str(cell).strip()) for cell in row] for row in table.extract()]
            except Exception:
                rows = []
            table_caption = pending_table_captions.pop(0) if pending_table_captions else None
            caption_text = table_caption["text"] if table_caption else f"Imported Table {len(tables) + 1}"
            table_id = f"table-{len(tables) + 1}"
            table_bbox = bbox(table.bbox[0], table.bbox[1], table.bbox[2] - table.bbox[0], table.bbox[3] - table.bbox[1])
            tables.append({"table_id": table_id, "page_number": page_index + 1, "title": caption_text, "caption": caption_text, "rows": rows, "row_count": len(rows), "column_count": len(rows[0]) if rows and rows[0] else 0, "geometry": table_bbox})
            layout_semantics.append({"kind": "table", "page_number": page_index + 1, "bbox": table_bbox, "row_count": len(rows), "column_count": len(rows[0]) if rows and rows[0] else 0})
            geometry_map.append({
                "geometry_id": f"pdf-geometry-{len(geometry_map) + 1}",
                "page_number": page_index + 1,
                "order_index": len(text_blocks) + table_index + 1,
                "kind": "table",
                "bbox": table_bbox,
                "table_id": table_id,
                "caption": caption_text
            })
            ensure_section(current_section["title"] if current_section else "Imported Tables")
            if preface_blocks:
                current_section["blocks"].extend(preface_blocks)
                preface_blocks = []
            current_section["blocks"].append({"block_type": "table", "title": caption_text, "body": caption_text, "dataset_ref": None, "query_ref": None, "field_mappings": [], "citations": [table_caption["caption_id"]] if table_caption else [], "table_rows": rows, "chart_series": [], "source_metadata": {"page_number": page_index + 1, "bbox": table_bbox, "section_heading_level": current_heading_level}, "layout_semantics": {"page_number": page_index + 1, "kind": "table", "bbox": table_bbox, "page_semantics_role": "table_region"}, "caption": caption_text, "page_number": page_index + 1, "source_lineage_refs": [f"pdf://page/{page_index + 1}/table/{table_index + 1}"]})
            block_lineage_map.append({"source_ref": f"pdf://page/{page_index + 1}/table/{table_index + 1}", "target_kind": "table", "target_ref": table_id, "caption_ref": table_caption["caption_id"] if table_caption else None, "page_number": page_index + 1})
        for drawing_index, drawing in enumerate(drawings):
            rect = drawing.get("rect")
            if rect is not None:
                vector_ref = f"pdf://page/{page_index + 1}/vector/{drawing_index + 1}"
                vector_bbox = bbox(rect.x0, rect.y0, rect.width, rect.height)
                embedded_assets.append({"asset_ref": vector_ref, "asset_kind": "vector_region", "page_number": page_index + 1, "bbox": vector_bbox, "path_count": len(drawing.get("items", []))})
                geometry_map.append({"geometry_id": f"pdf-geometry-{len(geometry_map) + 1}", "page_number": page_index + 1, "order_index": len(text_blocks) + len(found_tables) + drawing_index + 1, "kind": "vector_region", "bbox": vector_bbox, "vector_ref": vector_ref, "path_count": len(drawing.get("items", []))})
        for image_index, image_info in enumerate(page_images):
            if image_info.get("bbox") is not None:
                img_bbox = image_info["bbox"]
                embedded_assets.append({"asset_ref": f"pdf://page/{page_index + 1}/image/{image_index + 1}", "asset_kind": "image_region", "page_number": page_index + 1, "bbox": bbox(img_bbox[0], img_bbox[1], img_bbox[2] - img_bbox[0], img_bbox[3] - img_bbox[1])})
        if pending_caption and (len(drawings) > 0 or image_count > 0):
            series_seed = pending_caption.get("series_values", [])[:6] or [v for v in [len(drawings), image_count] if v > 0]
            chart_series = [{"label": f"Signal {i + 1}", "value": v} for i, v in enumerate(series_seed)]
            chart_id = f"chart-{len(charts) + 1}"
            visual_bbox = bbox(24, 24, 547, 794)
            charts.append({"chart_id": chart_id, "page_number": page_index + 1, "title": pending_caption["text"], "caption": pending_caption["text"], "chart_type": "bar", "series": chart_series, "source_anchor": f"page-{page_index + 1}", "geometry": visual_bbox})
            geometry_map.append({
                "geometry_id": f"pdf-geometry-{len(geometry_map) + 1}",
                "page_number": page_index + 1,
                "order_index": len(text_blocks) + len(found_tables) + 1,
                "kind": "visual_region",
                "bbox": visual_bbox,
                "chart_id": chart_id,
                "caption": pending_caption["text"]
            })
            ensure_section(current_section["title"] if current_section else "Imported Visuals")
            current_section["blocks"].append({"block_type": "chart", "title": pending_caption["text"], "body": pending_caption["text"], "dataset_ref": None, "query_ref": None, "field_mappings": [], "citations": [pending_caption["caption_id"]], "table_rows": [], "chart_series": chart_series, "source_metadata": {"page_number": page_index + 1, "drawings_detected": len(drawings), "images_detected": image_count, "section_heading_level": current_heading_level}, "layout_semantics": {"page_number": page_index + 1, "kind": "visual_region", "bbox": visual_bbox, "page_semantics_role": "visual_chart_region"}, "caption": pending_caption["text"], "page_number": page_index + 1, "source_lineage_refs": [f"pdf://page/{page_index + 1}/visual/{len(charts)}"]})
            block_lineage_map.append({"source_ref": f"pdf://page/{page_index + 1}/visual/{len(charts)}", "target_kind": "chart", "target_ref": chart_id, "caption_ref": pending_caption["caption_id"], "page_number": page_index + 1})
            pending_caption = None

    section_blueprints = [section for section in section_blueprints if section["blocks"]]
    if preface_blocks and section_blueprints:
        section_blueprints[0]["blocks"] = [*preface_blocks, *section_blueprints[0]["blocks"]]
    if not section_blueprints:
        ensure_section(title or Path(input_path).stem)
        if preface_blocks:
            current_section["blocks"].extend(preface_blocks)
        else:
            fallback_ref = f"pdf://fallback/{Path(input_path).stem}"
            current_section["blocks"].append({"block_type": "narrative", "title": title or Path(input_path).stem, "body": "\n".join(paragraphs) or Path(input_path).stem, "dataset_ref": None, "query_ref": None, "field_mappings": [], "citations": [], "table_rows": [], "chart_series": [], "source_metadata": {"page_number": 1, "fallback": True}, "layout_semantics": {"page_number": 1, "kind": "fallback", "bbox": bbox(24, 24, 547, 794)}, "caption": "", "page_number": 1, "source_lineage_refs": [fallback_ref]})
            block_lineage_map.append({"source_ref": fallback_ref, "target_kind": "paragraph", "target_ref": fallback_ref, "page_number": 1, "section_title": title or Path(input_path).stem})
    warnings = []
    if not paragraphs:
        warnings.append("pdf_text_extract_sparse")
    if not tables:
        warnings.append("pdf_table_extract_sparse")
    return {"title": title or Path(input_path).stem, "paragraphs": paragraphs, "extracted_text": "\n".join(extracted_lines), "warning_codes": warnings, "page_count": doc.page_count, "section_count": len(section_blueprints), "table_count": len(tables), "chart_count": len(charts), "caption_count": len(captions), "captions": captions, "tables": tables, "charts": charts, "page_structure": page_structure, "layout_semantics": layout_semantics, "geometry_map": geometry_map, "page_semantics": page_semantics, "section_hierarchy": section_hierarchy, "block_lineage_map": block_lineage_map, "section_blueprints": section_blueprints, "rendered_pages": rendered_pages, "source_language": infer_lang("\n".join(extracted_lines)), "hyperlinks": hyperlinks, "embedded_assets": embedded_assets, "document_metadata": {"page_count": doc.page_count, "metadata": doc.metadata}}

def build_sample_docx(output_path):
    doc = Document()
    doc.add_heading("Professional Imported Report", level=0)
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph("Coverage remained stable across the imported workstream and the editable report must preserve layout semantics.")
    doc.add_paragraph("Figure 1: Coverage Trend 94 91 88 86")
    doc.add_paragraph("The trend remained visible across the imported pages.")
    doc.add_paragraph("Table 1: KPI Coverage")
    table = doc.add_table(rows=1, cols=3)
    for i, value in enumerate(["KPI", "Current", "Target"]):
        table.rows[0].cells[i].text = value
    for row_values in [("Coverage", "94", "90"), ("Response Time", "18", "20"), ("Aging", "12", "15")]:
        row = table.add_row().cells
        for i, value in enumerate(row_values):
            row[i].text = value
    doc.add_page_break()
    doc.add_heading("الملخص العربي", level=1)
    doc.add_paragraph("يحافظ الاستيراد عالي الدقة على العناوين والجداول والتعليقات التوضيحية وبنية الصفحات.")
    doc.add_paragraph("جدول 2: متابعة المخاطر")
    table2 = doc.add_table(rows=1, cols=3)
    for i, value in enumerate(["الخطر", "المالك", "الحالة"]):
        table2.rows[0].cells[i].text = value
    for row_values in [("تأخير الاعتماد", "PMO", "قيد المعالجة"), ("تأخر المورد", "Ops", "تصعيد")]:
        row = table2.add_row().cells
        for i, value in enumerate(row_values):
            row[i].text = value
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return {
        "output_path": output_path,
        "fixture_profile": "basic",
        "expected": {"page_count": 2, "section_count": 2, "table_count": 2, "chart_count": 1, "caption_count": 3}
    }

def build_complex_sample_docx(output_path):
    doc = Document()
    doc.core_properties.title = "Complex Imported DOCX Report"
    doc.core_properties.author = "Codex"
    doc.core_properties.subject = "Complex DOCX fidelity fixture"
    doc.core_properties.comments = "Includes hyperlinks, media, section hierarchy, and mixed language content."
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.header.paragraphs[0].text = "Rasid Complex Header"
    section.footer.paragraphs[0].text = "Confidential / Editable"

    title = doc.add_heading("Complex Professional Report", level=0)
    title.alignment = 1
    intro = doc.add_paragraph("Executive Summary and source links are preserved in the editable artifact.")
    add_hyperlink(intro, " Source reference", "https://example.com/rasid/source")
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph("Figure 1: Coverage trend 98 94 91 88")
    doc.add_paragraph("The first page contains mixed English and Arabic narrative with hyperlink preservation.")
    chart_image = write_sample_png(Path(output_path).parent / ".assets")
    doc.add_picture(str(chart_image), width=Inches(5.6))
    doc.add_paragraph("شكل 2: مؤشرات الاستقرار 88 91 96")
    doc.add_paragraph("يحافظ الاستيراد على الصورة المضمّنة والتعليق التوضيحي والمراجع الخارجية.")
    doc.add_paragraph("Table 1: KPI Coverage")
    table = doc.add_table(rows=1, cols=4)
    for i, value in enumerate(["KPI", "Current", "Target", "Owner"]):
        table.rows[0].cells[i].text = value
    for row_values in [("Coverage", "98", "95", "PMO"), ("Aging", "12", "15", "Ops"), ("Escalations", "3", "5", "QA")]:
        row = table.add_row().cells
        for i, value in enumerate(row_values):
            row[i].text = value

    doc.add_page_break()
    landscape = doc.add_section(WD_SECTION_START.NEW_PAGE)
    landscape.orientation = WD_ORIENT.LANDSCAPE
    landscape.page_width, landscape.page_height = landscape.page_height, landscape.page_width
    landscape.header.paragraphs[0].text = "Landscape appendix header"
    landscape.footer.paragraphs[0].text = "Appendix footer"
    doc.add_heading("الملخص العربي", level=1)
    doc.add_paragraph("يحافظ المسار عالي الدقة على العناوين والصفحات والجداول والرسوم والروابط داخل المستند المعقد.")
    appendix_link = doc.add_paragraph("انظر الملحق المرجعي")
    add_hyperlink(appendix_link, " هنا", "https://example.com/rasid/appendix")
    doc.add_heading("القسم 2.1: هيكل الجداول", level=2)
    doc.add_paragraph("جدول 2: متابعة المخاطر")
    table2 = doc.add_table(rows=1, cols=4)
    for i, value in enumerate(["الخطر", "المالك", "الحالة", "التأثير"]):
        table2.rows[0].cells[i].text = value
    for row_values in [("تأخير الاعتماد", "PMO", "قيد المعالجة", "عال"), ("تأخر المورد", "Ops", "تصعيد", "متوسط"), ("انحراف البيانات", "Data", "مفتوح", "عال")]:
        row = table2.add_row().cells
        for i, value in enumerate(row_values):
            row[i].text = value
    doc.add_heading("Appendix", level=1)
    doc.add_paragraph("Figure 3: Recovery curve 74 82 91 97")
    doc.add_paragraph("The appendix mixes section hierarchy, captions, and page semantics for deeper parsing.")
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return {
        "output_path": output_path,
        "fixture_profile": "complex",
        "expected": {"page_count": 2, "section_count": 4, "table_count": 2, "chart_count": 3, "caption_count": 5, "hyperlink_count": 2, "embedded_asset_count": 2}
    }

def build_sample_pdf(output_path):
    font_name = ar_font()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("ReportTitle", parent=styles["Title"], fontName=font_name, fontSize=18, leading=24, alignment=TA_RIGHT)
    body_style = ParagraphStyle("ReportBody", parent=styles["BodyText"], fontName=font_name, fontSize=11, leading=16, alignment=TA_RIGHT)
    caption_style = ParagraphStyle("Caption", parent=styles["BodyText"], fontName=font_name, fontSize=10, textColor=colors.HexColor("#4B5563"), leading=14, alignment=TA_RIGHT)
    doc = SimpleDocTemplate(output_path, pagesize=A4, rightMargin=1.8 * cm, leftMargin=1.8 * cm, topMargin=1.6 * cm, bottomMargin=1.6 * cm)
    story = [
        PdfParagraph(shape("Professional Imported PDF Report", True), title_style),
        Spacer(1, 0.4 * cm),
        PdfParagraph(shape("الملخص التنفيذي", True), title_style),
        PdfParagraph(shape("يحافظ الاستيراد الأقوى على بنية الصفحة والعناوين والجداول والتسميات التوضيحية حتى في PDF غير اللاتيني.", True), body_style),
        Spacer(1, 0.2 * cm),
        PdfParagraph(shape("Figure 1: اتجاه التغطية 94 91 88 86", True), caption_style),
        PdfParagraph(shape("تمثل الأرقام أعلاه مسار المؤشر الرئيسي عبر أربع نقاط زمنية.", True), body_style),
        Spacer(1, 0.2 * cm),
        PdfParagraph(shape("Table 1: KPI Coverage", True), caption_style),
    ]
    table = PdfTable([["KPI", "Current", "Target"], ["Coverage", "94", "90"], ["Response", "18", "20"], ["Aging", "12", "15"]], colWidths=[6 * cm, 4 * cm, 4 * cm], hAlign="RIGHT")
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#94A3B8")), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E2E8F0")), ("FONTNAME", (0, 0), (-1, -1), font_name), ("FONTSIZE", (0, 0), (-1, -1), 10)]))
    story.extend([table, Spacer(1, 0.5 * cm), PdfParagraph(shape("جدول 2: متابعة المخاطر", True), caption_style)])
    table2 = PdfTable([["الخطر", "المالك", "الحالة"], ["تأخير الاعتماد", "PMO", "قيد المعالجة"], ["تأخر المورد", "Ops", "تصعيد"]], colWidths=[6 * cm, 4 * cm, 4 * cm], hAlign="RIGHT")
    table2.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#94A3B8")), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E2E8F0")), ("FONTNAME", (0, 0), (-1, -1), font_name), ("FONTSIZE", (0, 0), (-1, -1), 10)]))
    story.append(table2)
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.build(story)
    return {
        "output_path": output_path,
        "fixture_profile": "basic",
        "expected": {"page_count": 1, "section_count": 1, "table_count": 2, "chart_count": 1, "caption_count": 3}
    }

def build_complex_sample_pdf(output_path):
    font_name = ar_font()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("ComplexTitle", parent=styles["Title"], fontName=font_name, fontSize=18, leading=24, alignment=TA_RIGHT)
    heading_style = ParagraphStyle("ComplexHeading", parent=styles["Heading2"], fontName=font_name, fontSize=15, leading=20, alignment=TA_RIGHT)
    body_style = ParagraphStyle("ComplexBody", parent=styles["BodyText"], fontName=font_name, fontSize=11, leading=16, alignment=TA_RIGHT)
    caption_style = ParagraphStyle("ComplexCaption", parent=styles["BodyText"], fontName=font_name, fontSize=10, textColor=colors.HexColor("#4B5563"), leading=14, alignment=TA_RIGHT)

    def draw_page(canvas, _doc):
        canvas.setFont(font_name, 10)
        canvas.drawRightString(560, 820, shape("Rasid external complex PDF", True))
        canvas.drawRightString(560, 22, shape("Page %s" % canvas.getPageNumber(), True))
        canvas.linkURL("https://example.com/rasid/pdf-gateway", (420, 800, 560, 820), relative=0)

    story = [
        PdfParagraph(shape("Complex Imported PDF Report", True), title_style),
        Spacer(1, 0.3 * cm),
        PdfParagraph(shape("Executive Summary", True), heading_style),
        PdfParagraph(shape("يحافظ المسار العميق على البنية والاتجاه والرسوم المتجهية والروابط داخل ملفات PDF المعقدة.", True), body_style),
        Spacer(1, 0.2 * cm),
        PdfParagraph(shape("Figure 1: اتجاه التغطية 99 96 93 90", True), caption_style),
        PdfParagraph(shape("التحليل السردي مرتبط بمنطقة رسم متجهي مرئية في الصفحة نفسها.", True), body_style),
        Spacer(1, 0.2 * cm),
        PdfParagraph(shape("Table 1: KPI Coverage", True), caption_style)
    ]
    table = PdfTable([["KPI", "Current", "Target", "Owner"], ["Coverage", "99", "95", "PMO"], ["Response", "16", "20", "Ops"], ["Backlog", "7", "10", "QA"]], colWidths=[4.2 * cm, 3.4 * cm, 3.4 * cm, 3.4 * cm], hAlign="RIGHT")
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#94A3B8")), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E2E8F0")), ("FONTNAME", (0, 0), (-1, -1), font_name), ("FONTSIZE", (0, 0), (-1, -1), 10)]))
    story.extend([table, Spacer(1, 0.4 * cm)])
    drawing = Drawing(420, 180)
    drawing.add(Rect(20, 25, 360, 120, strokeColor=colors.HexColor("#0D5B56"), fillColor=colors.HexColor("#E6FFFA")))
    drawing.add(Line(45, 45, 45, 135, strokeColor=colors.HexColor("#1F2937"), strokeWidth=1))
    drawing.add(Line(45, 45, 360, 45, strokeColor=colors.HexColor("#1F2937"), strokeWidth=1))
    for index, value in enumerate([90, 93, 96, 99]):
        x = 80 + index * 65
        drawing.add(Rect(x, 45, 36, value - 60, strokeColor=colors.HexColor("#BA5D2A"), fillColor=colors.HexColor("#F59E0B")))
        drawing.add(String(x - 2, 25, f"Q{index + 1}", fontName=font_name, fontSize=10))
        drawing.add(String(x, value - 10, str(value), fontName=font_name, fontSize=10))
    drawing.add(String(230, 160, shape("Vector Coverage Chart", True), fontName=font_name, fontSize=12))
    story.extend([drawing, Spacer(1, 0.2 * cm), PdfParagraph(shape("شكل 2: استقرار الخدمات 84 88 91 95", True), caption_style), PageBreak()])
    story.extend([PdfParagraph(shape("الملحق العربي", True), heading_style), PdfParagraph(shape("يعرض هذا الملحق بنية صفحة متعددة وروابط وتخطيطًا معقدًا للمقارنة والاسترجاع.", True), body_style), PdfParagraph(shape("جدول 2: متابعة المخاطر", True), caption_style)])
    table2 = PdfTable([["الخطر", "المالك", "الحالة", "التأثير"], ["تأخير الاعتماد", "PMO", "قيد المعالجة", "عال"], ["تأخر المورد", "Ops", "تصعيد", "متوسط"], ["انحراف البيانات", "Data", "مفتوح", "عال"]], colWidths=[4.8 * cm, 3.2 * cm, 3.4 * cm, 2.8 * cm], hAlign="RIGHT")
    table2.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#94A3B8")), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E2E8F0")), ("FONTNAME", (0, 0), (-1, -1), font_name), ("FONTSIZE", (0, 0), (-1, -1), 10)]))
    story.append(table2)
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(output_path, pagesize=A4, rightMargin=1.4 * cm, leftMargin=1.4 * cm, topMargin=1.4 * cm, bottomMargin=1.4 * cm)
    doc.build(story, onFirstPage=draw_page, onLaterPages=draw_page)
    pdf = fitz.open(output_path)
    for page_index, page in enumerate(pdf):
        page.insert_link({"kind": fitz.LINK_URI, "from": fitz.Rect(420, 800, 560, 820), "uri": f"https://example.com/rasid/pdf-proof/page-{page_index + 1}"})
    pdf.save(output_path, incremental=True, encryption=fitz.PDF_ENCRYPT_KEEP)
    pdf.close()
    return {
        "output_path": output_path,
        "fixture_profile": "complex",
        "expected": {"page_count": 2, "section_count": 2, "table_count": 2, "chart_count": 2, "caption_count": 4, "hyperlink_count": 4, "vector_region_count": 1}
    }

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("command")
    parser.add_argument("input")
    parser.add_argument("output")
    parser.add_argument("--render-dir", default=None)
    args = parser.parse_args()
    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    if args.command == "extract-docx":
        payload = extract_docx(args.input)
    elif args.command == "extract-pdf":
        payload = extract_pdf(args.input, args.render_dir)
    elif args.command == "build-sample-docx":
        payload = build_sample_docx(args.input)
    elif args.command == "build-sample-pdf":
        payload = build_sample_pdf(args.input)
    elif args.command == "build-complex-sample-docx":
        payload = build_complex_sample_docx(args.input)
    elif args.command == "build-complex-sample-pdf":
        payload = build_complex_sample_pdf(args.input)
    else:
        raise SystemExit(f"Unknown command: {args.command}")
    output.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

if __name__ == "__main__":
    main()
